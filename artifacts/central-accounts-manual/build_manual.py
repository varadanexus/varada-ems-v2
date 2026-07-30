from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Varada EMS 2.0")
OUT = ROOT / "artifacts" / "central-accounts-manual"
SCREENS = OUT / "screens"
FINAL = OUT / "Varada_Nexus_Central_Accounts_Training_Manual.docx"
LOGO = ROOT / "new-ems" / "assets" / "pdf" / "vn-logo.png"

INK = "161616"
GOLD = "C8A44D"
DARK_GOLD = "7B5B19"
MUTED = "666666"
PALE_GOLD = "F7F1E2"
PALE_GRAY = "F3F4F6"
GREEN = "1F7A4D"
RED = "9B1C1C"
BLUE = "315B7D"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=None, color=INK, italic=None, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 17, DARK_GOLD, 16, 8),
        ("Heading 2", 13.5, INK, 12, 6),
        ("Heading 3", 11.5, DARK_GOLD, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.68)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.32)
        section.footer_distance = Inches(0.32)


def add_running_furniture(section):
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.86))
    set_table_geometry(table, [5200, 4678], indent_dxa=0)
    table._tbl.tblPr.append(OxmlElement("w:tblBorders"))
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = left.add_run("VARADA NEXUS PRIVATE LIMITED")
    set_font(r, size=8, bold=True, color=DARK_GOLD)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("CENTRAL ACCOUNTS TRAINING MANUAL")
    set_font(r, size=8, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Internal Training • Version 1.0 • Page ")
    set_font(r, size=8.5, color=MUTED)
    add_page_number(p)


def add_title(doc, text, size=28, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=color)
    return p


def add_kicker(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_font(r, size=9, bold=True, color=GOLD)
    r.font.all_caps = True
    return p


def add_para(doc, text, bold_lead=None, italic=False, align=None, after=6, color=INK):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_font(r)
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, title, text, tone="gold"):
    fill = {"gold": PALE_GOLD, "gray": PALE_GRAY, "green": "E7F5ED", "red": "FCECEC"}.get(tone, PALE_GOLD)
    accent = {"gold": DARK_GOLD, "gray": MUTED, "green": GREEN, "red": RED}.get(tone, DARK_GOLD)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9878], indent_dxa=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_dxa, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for idx, header in enumerate(headers):
        set_cell_shading(hdr.cells[idx], INK)
        p = hdr.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_font(r, size=font_size, bold=True, color=WHITE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if row_idx % 2:
                set_cell_shading(cells[idx], "FAFAFA")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, filename, caption, intro, notes):
    path = SCREENS / filename
    with Image.open(path) as source:
        crop_right = min(source.width, 1180)
        cropped = source.crop((0, 0, crop_right, source.height))
        cropped_path = SCREENS / f"cropped-{filename}"
        cropped.save(cropped_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.add_run().add_picture(str(cropped_path), width=Inches(4.85))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    r = cap.add_run(caption)
    set_font(r, size=8.5, bold=True, color=MUTED, italic=True)
    add_para(doc, intro, after=4)
    for note in notes:
        add_bullet(doc, note)


def make_flow_diagram():
    path = OUT / "central-accounts-flow.png"
    width, height = 1800, 410
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 42)
        box_font = ImageFont.truetype("arialbd.ttf", 27)
        sub_font = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        title_font = box_font = sub_font = ImageFont.load_default()
    draw.text((70, 35), "CENTRAL ACCOUNTS CONTROL FLOW", fill=(22, 22, 22), font=title_font)
    steps = [
        ("1", "SOURCE", "Approved division\ndocument"),
        ("2", "STAGE", "Financial\ndocument"),
        ("3", "VALIDATE", "Posting queue\nchecks"),
        ("4", "POST", "Balanced journal\nentry"),
        ("5", "RECONCILE", "AR • AP • Bank\nGST • TDS"),
        ("6", "CLOSE", "Review, evidence\nand lock"),
        ("7", "REPORT", "Statements and\nmanagement reports"),
    ]
    x = 55
    box_w = 220
    gap = 25
    y = 135
    for idx, (num, title, sub) in enumerate(steps):
        draw.rounded_rectangle((x, y, x + box_w, y + 190), radius=22, fill=(247, 241, 226), outline=(200, 164, 77), width=4)
        draw.ellipse((x + 14, y + 14, x + 60, y + 60), fill=(22, 22, 22))
        draw.text((x + 31, y + 22), num, fill=(255, 255, 255), font=sub_font, anchor="mm")
        draw.text((x + 75, y + 25), title, fill=(123, 91, 25), font=box_font)
        for line_idx, line in enumerate(sub.split("\n")):
            draw.text((x + box_w / 2, y + 95 + line_idx * 30), line, fill=(40, 40, 40), font=sub_font, anchor="mm")
        if idx < len(steps) - 1:
            ax = x + box_w + 5
            ay = y + 95
            draw.line((ax, ay, ax + gap - 8, ay), fill=(123, 91, 25), width=5)
            draw.polygon([(ax + gap - 8, ay - 9), (ax + gap + 4, ay), (ax + gap - 8, ay + 9)], fill=(123, 91, 25))
        x += box_w + gap
    image.save(path)
    return path


def section_title(doc, number, title, subtitle=None):
    h = doc.add_heading(f"{number}. {title}", level=1)
    if subtitle:
        add_para(doc, subtitle, italic=True, color=MUTED, after=8)
    return h


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    flow_path = make_flow_diagram()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_running_furniture(doc.sections[0])

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.add_run().add_picture(str(LOGO), width=Inches(1.7))
    add_kicker(doc, "Internal Training & Operating Guide", WD_ALIGN_PARAGRAPH.CENTER)
    add_title(doc, "Central Accounts", size=31, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_title(doc, "Training Manual", size=23, color=DARK_GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=15)
    add_para(doc, "From source documents to postings, reconciliations, statutory controls, close and reporting", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, color=MUTED, after=26)
    add_callout(doc, "Training objective", "Enable a new finance user to understand what each Central Accounts screen does, complete routine work safely, recognize when approval is required, and preserve a complete audit trail.", "gold")
    add_para(doc, "Prepared for: Accounts Executives, Accounts Managers, CFO/CA reviewers, authorised auditors and system administrators", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, after=5)
    add_para(doc, "Version 1.0 • 20 July 2026 • Varada Nexus Private Limited", align=WD_ALIGN_PARAGRAPH.CENTER, bold_lead="Version 1.0", after=4)
    add_para(doc, "System reference: EMS 2.0 release/2.0-rc1", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    doc.add_page_break()

    section_title(doc, "1", "How to use this manual")
    add_para(doc, "Read sections 1–6 before posting any transaction. Use sections 7–15 as operating references for the area assigned to you. Complete the closing checklist in section 16 at month-end.")
    add_callout(doc, "Important control", "Central Accounts is a controlled ledger workspace. A document should be posted only after the source is approved, the accounting treatment is verified, the period is open, supporting evidence is available and the debit/credit effect is understood.", "red")
    doc.add_heading("Learning outcomes", level=2)
    for item in [
        "Navigate all Central Accounts workspaces and understand what belongs in each one.",
        "Trace an approved source transaction into a financial document, posting queue item and journal.",
        "Maintain receivables, payables, treasury, tax, fixed asset and budget records without bypassing approvals.",
        "Prepare close evidence and reliable reports while keeping the audit trail intact.",
        "Recognize errors that must be corrected before posting instead of being hidden after posting.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("The three rules to remember", level=2)
    add_table(doc, ["Rule", "Meaning in daily work"], [
        ["Source first", "Every posting must be supported by an approved source document, authorised voucher or properly documented adjustment."],
        ["Maker and checker", "The person preparing data should not silently approve or post their own work when separate review is required."],
        ["Never erase history", "Correct mistakes through the approved correction, reversal or adjustment route; do not attempt to hide the original event."],
    ], [1900, 7478])

    section_title(doc, "2", "The complete accounting flow", "Understand the sequence before learning individual screens.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(flow_path), width=Inches(6.7))
    add_para(doc, "Figure 1 — Central Accounts control flow", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, color=MUTED, after=8)
    for text in [
        "Source: A division approves a bill, invoice, credit note, receipt, transporter statement, payment, expense or other eligible transaction.",
        "Stage: EMS creates a Central Accounts financial document without yet changing the final journal ledger.",
        "Validate: Finance checks document family, amount, date, counterparty, tax treatment, dimensions and queue status.",
        "Post: An authorised user confirms posting. EMS creates a balanced journal and updates relevant open items.",
        "Reconcile: Finance compares receivables, payables, bank movements, GST, TDS and division balances to supporting records.",
        "Close and report: Review evidence is completed, the period is controlled or locked, and reports are issued from posted records.",
    ]:
        add_step(doc, text)
    add_callout(doc, "What posting means", "Posting is the moment the transaction becomes part of the accounting books. Review before you press Post. A successful toast and changed queue status are the expected system confirmation.", "green")

    section_title(doc, "3", "Access, roles and navigation")
    add_figure(doc, "01-dashboard.png", "Figure 2 — Central Accounts Dashboard and sidebar", "The dashboard is the starting point. The left sidebar groups the module by the job being performed, while the control hub shows current posting health.", [
        "Home: Dashboard and Audit Events.",
        "Accounting Operations: Financial Documents, Posting Queue, Journals and Manual Vouchers.",
        "Working Books: Receivables, Payables, Treasury and Fixed Assets.",
        "Statutory & Controls: GST, TDS, Annual Tax, Close Controls and Tax & Company Settings.",
        "Financial Reporting: Reporting, Consolidated Books and Budgets & Profitability.",
    ])
    doc.add_heading("Typical responsibilities", level=2)
    add_table(doc, ["User type", "Typical work", "Control expectation"], [
        ["Accounts Executive", "Prepare masters, records, reconciliations and supporting evidence.", "Normally creates/edits; posting and approval depend on assigned permissions."],
        ["Accounts Manager", "Reviews treatment, approves controlled items, posts where authorised and supervises close.", "Checks source, tax, dimensions and balances before approval/posting."],
        ["CFO / CA reviewer", "Reviews financial statements, statutory positions, major adjustments and close readiness.", "Challenges exceptions and approves only with sufficient evidence."],
        ["Auditor", "Examines records, evidence, events and reports according to granted rights.", "Must not alter operational records unless explicitly authorised."],
        ["Administrator", "Maintains access and technical configuration.", "System access does not replace accounting approval."],
    ], [1650, 4020, 3708], font_size=8.8)
    add_callout(doc, "Permission-aware interface", "Buttons appear only when your role has the required View, Create, Edit, Approve, Post, Reconcile or Export action. If a button is missing, do not share another user’s login; ask an administrator to review your assigned permissions.", "gray")

    section_title(doc, "4", "Dashboard: start-of-day control check")
    add_para(doc, "At the start of each work session, use the dashboard to decide where attention is needed. The figures are control signals, not a replacement for ledger review.")
    add_table(doc, ["Indicator", "What it tells you", "Action"], [
        ["Ready To Post", "Validated documents waiting for authorised posting.", "Open Posting Queue; review oldest and highest-risk items first."],
        ["Posted", "Documents successfully posted to the ledger.", "Use Journals and reports to verify the resulting entry."],
        ["Failed", "Posting attempts that did not complete.", "Read the error, correct the underlying issue, and retry only after review."],
        ["Financial Documents", "Total staged/posted accounting documents in scope.", "Use filters to locate a source item and inspect its details."],
        ["Receivables / Payables", "Current open customer/vendor exposure shown by the dashboard dataset.", "Reconcile against working books and source division records."],
    ], [1700, 4250, 3428], font_size=8.8)
    add_callout(doc, "Do not rely on colour alone", "Always read the status text and amount. A zero balance may mean there are no posted open items, not that every source division has completed its billing workflow.", "gold")

    section_title(doc, "5", "Financial Documents: verify the staged source")
    add_figure(doc, "02-financial-documents.png", "Figure 3 — Financial Documents register", "This register is the bridge between approved operational documents and the accounting ledger. Use it to confirm what EMS intends to post.", [
        "Search by document number or family and filter by status.",
        "Review document family, date, counterparty/dimension and net amount.",
        "Open View to inspect queue context, posting sequence and the full record payload.",
        "Do not post from this screen; posting occurs in Posting Queue after review.",
    ])
    doc.add_heading("Document families you may see", level=2)
    add_table(doc, ["Family", "Business meaning", "Likely accounting effect"], [
        ["CLIENT_BILL / GST_INVOICE", "Amount billed to a customer.", "Revenue, output tax and receivable."],
        ["CLIENT_RECEIPT", "Money collected from a customer.", "Bank/cash increases and receivable reduces."],
        ["CREDIT_NOTE", "Reduction or correction to customer billing.", "Revenue/tax and receivable adjustment."],
        ["TRANSPORTER_STATEMENT", "Approved transporter/vendor liability.", "Expense/input tax and payable."],
        ["TRANSPORTER_PAYMENT", "Settlement of a transporter/vendor balance.", "Payable reduces and bank/cash reduces."],
        ["INTERIOR_BILL / division invoice", "Approved bill from another EMS division.", "Configured revenue/tax/receivable treatment."],
    ], [2300, 3778, 3300], font_size=8.5)
    doc.add_heading("Pre-post validation", level=2)
    for item in [
        "The source document number exists and matches the approved division record.",
        "The document date and accounting period are correct and open.",
        "The counterparty and reporting division/dimension are correct.",
        "Taxable value, GST/TDS components and net amount agree with the source.",
        "Required evidence is attached or referenced and there is no unresolved exception.",
    ]:
        add_bullet(doc, item)

    section_title(doc, "6", "Posting Queue: create the accounting entry")
    add_figure(doc, "03-posting-queue.png", "Figure 4 — Posting Queue", "The Posting Queue shows documents that are ready, processing, posted or failed. Only users with Post permission receive an active Post button.", [
        "Ready To Post means the item is eligible for a posting attempt; it does not remove the reviewer’s responsibility.",
        "Posting Sequence is the traceable link used to connect the document, posting job and journal.",
        "Failed items require diagnosis. Repeated clicking is not a correction method.",
        "EMS asks for confirmation before executing the posting.",
    ])
    doc.add_heading("Safe posting procedure", level=2)
    for step in [
        "Filter to Ready To Post and locate the intended source document.",
        "Return to Financial Documents and verify its family, amount, date, counterparty and queue context.",
        "Confirm the source division has approved the transaction and no later cancellation or revision exists.",
        "Check the accounting period, tax treatment and expected debit/credit impact.",
        "Select Post once and confirm the system prompt.",
        "Wait for the success message and refreshed queue status.",
        "Open Journals and verify the resulting entry is balanced and carries the correct reference.",
    ]:
        add_step(doc, step)
    add_callout(doc, "If posting fails", "Record the exact error and source document. Check missing account mappings, closed periods, invalid status, duplicated posting, unbalanced lines or missing master data. Correct the cause, then retry once. Escalate repeated failures to the Accounts Manager/system administrator.", "red")

    section_title(doc, "7", "Journals and Manual Vouchers")
    add_figure(doc, "04-manual-vouchers.png", "Figure 5 — Manual Voucher entry and register", "Use Manual Vouchers only for authorised accounting entries that do not originate from a normal source workflow, such as accruals, provisions, prepaids, opening entries and documented adjustments.", [
        "Choose the voucher type and date, then select the relevant division or company scope.",
        "Provide a meaningful reference and narration that another reviewer can understand.",
        "Add at least two lines and keep total debit exactly equal to total credit.",
        "Submitted vouchers move through approval and posting according to permissions.",
    ])
    doc.add_heading("Voucher types", level=2)
    add_table(doc, ["Type", "Use it for"], [
        ["journal / adjustment", "General reclassification or supported correcting entry."],
        ["payment", "Authorised payment entry not already generated by a controlled source workflow."],
        ["receipt", "Authorised receipt entry not already generated by the receivable workflow."],
        ["contra", "Transfer between cash/bank accounts."],
        ["accrual / provision", "Recognising an expense or liability before final settlement."],
        ["prepaid", "Recording or releasing expenditure relating to a future period."],
        ["opening", "Controlled opening balances during implementation or approved migration."],
    ], [2300, 7078], font_size=9)
    add_callout(doc, "Journal register", "The Journals screen is read-only and provides header/line drill-down. A valid journal must have equal debit and credit totals. Use the journal number and source reference when investigating balances.", "gray")

    section_title(doc, "8", "Receivables and Payables")
    add_figure(doc, "05-payables.png", "Figure 6 — Payables and vendor accounting", "Payables combines vendor master setup, purchase bill entry, vendor advances, approvals/posting and the open-item working book. Receivables is a read-only customer collection working book derived from posted items.", [
        "Vendor setup: vendor code, legal name, GSTIN/PAN and payable account.",
        "Purchase bill: vendor, division, dates, expense/input-tax/payable/TDS accounts and tax breakup.",
        "Advance: vendor, date, amount and traceable reference.",
        "Working books: original amount, open amount, due date, status and aging buckets.",
    ])
    doc.add_heading("Purchase bill entry — field guidance", level=2)
    add_table(doc, ["Field", "How to complete it"], [
        ["Vendor / Bill No", "Select the correct vendor and enter the invoice number exactly as shown on the supplier document."],
        ["Bill Date / Due Date", "Use the supplier invoice date and contractual due date."],
        ["Expense Account", "Choose the account that represents the nature of the goods or service."],
        ["Input Tax Account", "Use only when GST input credit is eligible and supported."],
        ["Payable Account", "Select the vendor liability/control account configured for AP."],
        ["TDS Payable Account", "Use when tax is deducted and payable to the government."],
        ["Taxable / CGST / SGST / IGST", "Copy from the valid tax invoice and verify intra-state versus inter-state treatment."],
        ["TDS / Total", "Check the deduction and invoice total; ensure the settlement amount agrees with the bill terms."],
    ], [2400, 6978], font_size=8.8)
    doc.add_heading("Aging interpretation", level=2)
    add_para(doc, "Aging buckets group open balances by how long they are past the due/reference date: 0–30, 31–60, 61–90 and 90+ days. Use aging to prioritise collections, payment scheduling and exception follow-up.")
    add_callout(doc, "Reconciliation rule", "The working book must reconcile to the relevant control-account balance in the general ledger. Investigate timing, unposted documents, partial settlements, credit notes, advances and incorrect counterparty dimensions.", "gold")

    section_title(doc, "9", "Treasury and bank reconciliation")
    add_figure(doc, "06-treasury.png", "Figure 7 — Treasury workspace", "Treasury manages cash/bank accounts, statement imports and reconciliation. The aim is to explain every bank line with a ledger movement or a documented exception.", [
        "Use filters to focus on a treasury account, date range or reconciliation status.",
        "Review book balance, statement balance and difference before signing off.",
        "Match bank statement lines to ledger movements; do not force a match when amount/date/reference do not support it.",
        "Generate or retain the reconciliation certificate and reviewer evidence when the period is complete.",
    ])
    doc.add_heading("Bank reconciliation procedure", level=2)
    for step in [
        "Confirm the treasury account is the correct bank/cash account and its opening/book balance is reasonable.",
        "Import or review the bank statement for the exact period without duplicate lines.",
        "Match receipts, payments, bank charges, interest and transfers using amount, date, reference and counterparty.",
        "Record legitimate book adjustments through the correct controlled entry; never alter a bank statement line to make it fit.",
        "Classify unmatched items and assign follow-up ownership.",
        "Verify the statement balance, book balance and reconciliation difference.",
        "Complete review evidence and obtain the required approval/reconciliation sign-off.",
    ]:
        add_step(doc, step)

    section_title(doc, "10", "GST compliance and company tax settings")
    add_figure(doc, "07-gst-compliance.png", "Figure 8 — GST Compliance", "GST Compliance brings together registration periods, classified documents, exception handling, GSTR-2B matching and input-tax-credit review. The system does not guess missing tax classifications.", [
        "Select the correct GST registration and return period.",
        "Review every classification or reconciliation exception.",
        "Check taxable value and CGST/SGST/IGST/cess against the source invoice.",
        "Review GSTR-2B match status and whether ITC is available or blocked.",
        "Lock books only after reconciliation and review are complete.",
    ])
    add_figure(doc, "09-tax-settings.png", "Figure 9 — Tax & Company Settings", "Tax & Company Settings is the master-data foundation for statutory workflows. It stores company identity, PAN/TAN/CIN, registered address, financial-year basis, auditor details and state-wise GST registrations.", [
        "Enter legal identifiers exactly as issued by the relevant authority.",
        "Maintain one row for each GST registration and mark only one appropriate active registration as primary.",
        "The page validates PAN and GSTIN formats but format validation is not proof that a registration is legally valid.",
        "Deactivation is controlled and audited; do not deactivate a registration with unresolved filing obligations.",
    ])
    doc.add_heading("GST review order", level=2)
    for item in [
        "Registration and filing period",
        "Outward invoices and credit notes",
        "Input documents and eligibility",
        "CGST/SGST versus IGST classification",
        "GSTR-2B match exceptions",
        "Tax ledgers, payment/offset and filing acknowledgement",
    ]:
        add_step(doc, item)
    add_callout(doc, "Professional review", "GST classification, ITC eligibility and return filing require review by the responsible Accounts Manager/CA. EMS supports preparation, reconciliation and evidence; it does not replace statutory judgement.", "red")

    section_title(doc, "11", "TDS and Annual Tax & Audit")
    doc.add_heading("TDS Compliance", level=2)
    add_para(doc, "Maintain deductees and TDS sections before recording deductions. For each deduction, link the deductee, applicable section, optional purchase bill, deduction date, taxable amount, rate, TDS amount, return period and challan reference.")
    add_table(doc, ["Control", "What to verify"], [
        ["Deductee identity", "Legal name, PAN, GSTIN where applicable and deductee type."],
        ["Section and rate", "Correct statutory section, threshold, default/lower deduction rate and supporting certificate."],
        ["Deduction timing", "Correct date of credit/payment and return quarter."],
        ["Challan linkage", "Payment challan number and amount agree with deductions included."],
        ["Return / certificate", "Quarterly return status, validation errors and certificates are tracked."],
    ], [2450, 6928], font_size=9)
    doc.add_heading("Annual Tax & Audit", level=2)
    add_para(doc, "Use the statutory filing calendar for GSTR-1, GSTR-3B, GSTR-9/9C, ITR, tax audit 3CD, TDS returns, advance tax and other obligations. Use Annual Audit Workpapers to record the financial year, section, title, amount, status and evidence reference.")
    add_callout(doc, "Evidence discipline", "Every filing marked complete should carry its acknowledgement/ARN and evidence link. Every significant annual-tax adjustment should be supported by a workpaper that another reviewer can reproduce.", "green")

    section_title(doc, "12", "Fixed Assets, budgets and profitability")
    doc.add_heading("Fixed Assets", level=2)
    add_para(doc, "The asset register captures asset code/name/class, owning division, acquisition and put-to-use dates, vendor/purchase bill, cost, salvage value, useful life and depreciation method. Movements record capitalization, transfer, impairment, revaluation, disposal or scrapping.")
    add_para(doc, "Depreciation runs progress from draft/prepared to approved and posted. Before posting, select the correct depreciation expense and accumulated depreciation accounts and confirm the run total agrees with the asset schedule.")
    doc.add_heading("Budgets & Profitability", level=2)
    add_para(doc, "Create a budget by financial year, division, scenario and status. Add lines by ledger account and period. Capture profitability snapshots to preserve a dated view of revenue, direct cost, indirect cost, gross margin and net margin.")
    add_table(doc, ["Scenario", "Use"], [
        ["base", "Approved operating expectation."],
        ["optimistic", "Upside scenario with documented assumptions."],
        ["conservative", "Downside scenario for risk planning."],
        ["reforecast", "Updated expectation based on actual performance and new information."],
    ], [2200, 7178], font_size=9.2)
    add_callout(doc, "Comparison control", "Actual profitability is sourced from posted/approved accounting records. A budget or snapshot does not change the ledger; it is a planning and analysis layer.", "gray")

    section_title(doc, "13", "Close Controls")
    add_figure(doc, "10-close-controls.png", "Figure 10 — Close Controls", "Close Controls turns month-end into an owned, evidenced checklist. Create one checklist per month, link the accounting period, assign tasks, track status and lock only after review.", [
        "Checklist statuses: open, in progress, ready for review, closed or reopened.",
        "Task statuses: pending, in progress, completed, reviewed, blocked or not applicable.",
        "Each task should have an area, owner, due date, status, evidence reference and useful notes.",
        "Lock Close and Reopen are approval-level actions; reopening must be justified and reviewed.",
    ])
    doc.add_heading("Minimum close areas", level=2)
    add_table(doc, ["Area", "Minimum completion evidence"], [
        ["Source completeness", "All divisions confirm approved bills, credit notes, receipts, statements and expenses are sent to Central Accounts."],
        ["Posting queue", "No unexplained ready, processing or failed items remain."],
        ["Receivables / Payables", "Open-item books reconcile to control accounts; aging exceptions have owners."],
        ["Bank / Treasury", "All bank accounts reconciled and certificates/evidence retained."],
        ["GST / TDS", "Registers reconciled, exceptions documented and filing/payment status updated."],
        ["Fixed assets", "Additions, disposals, transfers and depreciation reviewed and posted."],
        ["Journals", "Manual and automated journals reviewed for balance, narration, references and unusual entries."],
        ["Reports", "Trial balance and financial statements reviewed; material variances explained."],
    ], [2100, 7278], font_size=8.6)

    section_title(doc, "14", "Financial Reporting and Consolidated Books")
    add_figure(doc, "08-reporting.png", "Figure 11 — Central Accounts Reporting", "Reporting is built from existing journals, financial documents, open items and treasury movements. It does not create a duplicate ledger.", [
        "Choose the required period, division and report section before interpreting results.",
        "Use CSV for analysis, PDF for a fixed record and Print for controlled paper/save-to-PDF output.",
        "Check report totals against the journal/trial balance before external circulation.",
        "Keep generated date, filters and reviewer sign-off with formal reporting packs.",
    ])
    doc.add_heading("Reports and review questions", level=2)
    add_table(doc, ["Report", "Reviewer question"], [
        ["Trial Balance", "Do total debits equal total credits, and are unusual account balances explained?"],
        ["Profit & Loss", "Are revenue/cost classifications correct and material movements explained?"],
        ["Balance Sheet", "Do control accounts, taxes, bank, assets and equity reconcile to schedules?"],
        ["Receivables / Payables", "Do open items and aging agree with control accounts and source parties?"],
        ["Cash / Treasury", "Do book balances reconcile to bank statements and certificates?"],
        ["Division profitability", "Are reporting dimensions complete and allocations documented?"],
    ], [2400, 6978], font_size=8.8)
    add_para(doc, "Consolidated Books provides a company-wide document view across divisions, with document previews and downloadable PDFs where available. Use it to trace bills, receipts, payments and postings without switching repeatedly between source modules.")

    section_title(doc, "15", "Audit trail and control evidence")
    add_para(doc, "Audit Events is the investigation workspace. Search by event, actor, source document, financial document or journal reference. Each event records when it occurred, who performed it and the metadata available for that action.")
    add_para(doc, "The Auditor Workspace can maintain audit requests and maker-checker/control evidence. Requests include request number, year, area, title, due date, status and description. Control evidence records control code/area, period, entity, maker/checker actions and evidence status.")
    add_table(doc, ["Evidence status", "Meaning"], [
        ["captured", "Evidence has been recorded but may still require review."],
        ["exception", "A control failure, missing evidence or unexplained condition requires action."],
        ["reviewed", "An authorised reviewer has examined the evidence."],
        ["remediated", "The exception was corrected and the resolution is documented."],
    ], [2200, 7178], font_size=9.2)
    add_callout(doc, "Good narration test", "A reviewer unfamiliar with the transaction should be able to understand what happened, why it was authorised, which source supports it, who checked it and how the final balance was verified.", "gold")

    section_title(doc, "16", "Month-end operating checklist")
    checklist = [
        ("1", "Cut-off confirmed", "All divisions have submitted approved transactions for the period; late items are identified."),
        ("2", "Posting queue cleared", "Ready/processing/failed items are posted or documented with owner and resolution date."),
        ("3", "Journals reviewed", "Automated and manual entries are balanced, referenced and approved."),
        ("4", "AR/AP reconciled", "Working books agree with control accounts and aging exceptions are assigned."),
        ("5", "Bank reconciled", "All accounts reconcile; unmatched items and adjustments are evidenced."),
        ("6", "GST/TDS reconciled", "Registers, returns, challans, ITC and exceptions are reviewed."),
        ("7", "Assets completed", "Additions, disposals, transfers and depreciation are posted."),
        ("8", "Management review", "Trial balance, P&L, balance sheet and division results are reviewed."),
        ("9", "Evidence archived", "Close tasks contain links/references to support and reviewer sign-off."),
        ("10", "Period controlled", "Checklist is ready for review and lock/close is performed by an authorised approver."),
    ]
    add_table(doc, ["#", "Check", "Completion standard"], checklist, [600, 2350, 6428], font_size=8.8)
    add_callout(doc, "Reopening a close", "Reopen only when a material correction is required. Record the reason, responsible approver, affected entries and the checks repeated after correction.", "red")

    section_title(doc, "17", "Troubleshooting and common mistakes")
    add_table(doc, ["Problem", "Likely cause", "What to do"], [
        ["Post button missing", "No Post permission or item is not ready_to_post.", "Check role permissions and document/queue status; never borrow credentials."],
        ["Posting failed", "Missing mapping, closed period, invalid status, duplicate or unbalanced treatment.", "Read the error, correct the cause, retain evidence and retry once."],
        ["Document not in queue", "Source not approved/bridged, already posted or filtered out.", "Verify source workflow, clear filters and inspect Financial Documents."],
        ["AR/AP differs from ledger", "Unposted items, timing, credit notes, advances, partial settlements or wrong dimension.", "Reconcile item-by-item to source and journal references."],
        ["GST exception remains", "Tax breakup/classification/registration or 2B match is incomplete.", "Review source invoice and statutory eligibility; obtain CA review where needed."],
        ["Bank difference", "Unmatched statement item, duplicate import, missing book entry or timing difference.", "Investigate; record a valid adjustment only with approval and evidence."],
        ["Report looks incomplete", "Filters, period, division or unposted data.", "Confirm report parameters and posting completeness before exporting."],
        ["Button/action unavailable", "Page is permission-aware.", "Ask an administrator to grant the correct action to your role if business-authorised."],
    ], [2150, 3250, 3978], font_size=8.3)

    section_title(doc, "18", "Glossary")
    add_table(doc, ["Term", "Plain-language meaning"], [
        ["Financial document", "The Central Accounts representation of an approved source transaction."],
        ["Posting queue", "The controlled waiting area before a transaction changes the ledger."],
        ["Journal", "A balanced set of debit and credit lines recorded in the books."],
        ["Ledger account", "The account used to classify assets, liabilities, income, expenses or equity."],
        ["Reporting dimension", "A division, counterparty or other attribute used to analyse a posting."],
        ["Open item", "An unpaid or partly settled receivable/payable balance."],
        ["Aging", "Grouping open balances by the time elapsed from their due/reference date."],
        ["Reconciliation", "Proving that two related records or balances agree, with explanations for differences."],
        ["ITC", "Input Tax Credit, subject to statutory eligibility and evidence."],
        ["TDS", "Tax Deducted at Source, tracked by deductee, section, challan and return period."],
        ["Close", "The controlled completion and locking of an accounting period after review."],
        ["Maker-checker", "Separation between the preparer and the reviewer/approver of controlled work."],
    ], [2250, 7128], font_size=8.8)

    doc.add_heading("Training sign-off", level=2)
    add_para(doc, "Before receiving posting or approval rights, the trainee should demonstrate one supervised document review, one posting trace, one reconciliation walkthrough and one close-control task with evidence.")
    add_table(doc, ["Trainee", "Trainer", "Date", "Outcome"], [["", "", "", "Competent / Further training required"]], [2500, 2500, 1600, 2778], font_size=8.8)
    add_callout(doc, "Scope and disclaimer", "This manual explains how the EMS Central Accounts module is designed to be operated. It is not tax, audit, legal or accounting advice and does not override company policy, statutory requirements or professional judgement. Screen actions remain controlled by the user’s assigned permissions.", "gray")

    configure_page(doc)
    doc.core_properties.title = "Varada Nexus Central Accounts Training Manual"
    doc.core_properties.subject = "EMS 2.0 Central Accounts operator training"
    doc.core_properties.author = "Varada Nexus Private Limited"
    doc.core_properties.keywords = "Central Accounts, EMS, finance, accounting, training"
    doc.save(FINAL)
    print(FINAL)


if __name__ == "__main__":
    build()
